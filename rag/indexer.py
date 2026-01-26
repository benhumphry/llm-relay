"""
RAG Indexer service for Smart RAG.

Handles document indexing using Docling for parsing and ChromaDB for storage.
Supports scheduled indexing via APScheduler.
"""

import hashlib
import logging
import os
import threading
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, Union

logger = logging.getLogger(__name__)


# Supported file extensions (Docling capabilities)
SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf",
    ".docx",
    ".pptx",
    ".xlsx",
    ".html",
    ".htm",
    ".md",
    ".txt",
    ".asciidoc",
    ".adoc",
    # Images (OCR)
    ".png",
    ".jpg",
    ".jpeg",
    ".tiff",
    ".bmp",
}


class RAGIndexer:
    """
    Background service for indexing documents into ChromaDB.

    Uses Docling for document parsing and supports scheduled re-indexing.
    """

    def __init__(self):
        self._scheduler = None
        self._chroma_client = None
        self._running = False
        self._lock = threading.Lock()
        # Track indexing jobs by type and ID: "rag_{id}" or "store_{id}"
        self._indexing_jobs: dict[str, threading.Thread] = {}
        # Track cancelled jobs - checked by indexing loop to stop early
        self._cancelled_jobs: set[str] = set()

    def _clear_gpu_memory(self):
        """Clear GPU memory after indexing to prevent memory leaks."""
        try:
            import gc

            gc.collect()

            import torch

            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                torch.cuda.synchronize()
                logger.debug("Cleared CUDA memory cache")
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"Failed to clear GPU memory: {e}")

    def _get_unified_source_for_store(self, store: Any) -> Optional[Any]:
        """
        Check if a unified source plugin exists for this document store.

        Unified sources combine RAG + Live into a single plugin. When available,
        they provide the document enumeration and reading capabilities needed
        for indexing.

        Args:
            store: DocumentStore model instance

        Returns:
            Instantiated unified source plugin if available, None otherwise
        """
        source_type = getattr(store, "source_type", "")

        try:
            from plugin_base.loader import get_unified_source_for_doc_type

            # Dynamic lookup - find unified source that handles this doc store type
            plugin_class = get_unified_source_for_doc_type(source_type)
            if not plugin_class:
                return None

            # Use the plugin's own method to build config from store
            # This checks PluginConfig first, then falls back to legacy columns
            config = plugin_class.get_config_for_store(store)

            logger.info(
                f"Using unified source plugin '{plugin_class.source_type}' for store '{store.name}'"
            )
            return plugin_class(config)

        except Exception as e:
            logger.debug(f"Could not instantiate unified source for {source_type}: {e}")
            return None

    def _get_existing_doc_metadata(self, collection) -> dict[str, Optional[str]]:
        """
        Get existing document URIs and their modified times from a ChromaDB collection.

        Returns:
            Dict mapping source_uri -> modified_time (or None if not available)
        """
        existing_docs = {}
        try:
            # Get all documents with metadata
            results = collection.get(include=["metadatas"])
            if results and results.get("metadatas"):
                for meta in results["metadatas"]:
                    uri = meta.get("source_uri")
                    if uri and uri not in existing_docs:
                        existing_docs[uri] = meta.get("modified_time")
        except Exception as e:
            logger.warning(f"Failed to get existing documents: {e}")
        return existing_docs

    def _delete_doc_chunks(self, collection, source_uri: str) -> int:
        """
        Delete all chunks for a specific document URI from a collection.

        Returns:
            Number of chunks deleted
        """
        try:
            # Get IDs of chunks with this source_uri
            results = collection.get(where={"source_uri": source_uri}, include=[])
            if results and results.get("ids"):
                ids_to_delete = results["ids"]
                collection.delete(ids=ids_to_delete)
                logger.debug(f"Deleted {len(ids_to_delete)} chunks for {source_uri}")
                return len(ids_to_delete)
        except Exception as e:
            logger.warning(f"Failed to delete chunks for {source_uri}: {e}")
        return 0

    def generate_store_intelligence(
        self, store_id: int, intelligence_model: str | None = None
    ) -> bool:
        """
        Generate intelligence (themes, best_for, summary) for a document store.

        Samples chunks from the store and uses an LLM to analyze the content,
        generating themes, use cases, and a summary that helps routing decisions.

        Args:
            store_id: ID of the document store
            intelligence_model: Model to use for analysis (defaults to global setting)

        Returns:
            True if intelligence was generated successfully
        """
        from db import get_db_context, get_document_store_by_id
        from db.models import DocumentStore, Setting

        try:
            store = get_document_store_by_id(store_id)
            if not store:
                logger.warning(
                    f"Store {store_id} not found for intelligence generation"
                )
                return False

            if not store.collection_name or store.chunk_count == 0:
                logger.info(
                    f"Store '{store.name}' has no indexed content, skipping intelligence"
                )
                return False

            # Get the intelligence model from settings if not provided
            if not intelligence_model:
                with get_db_context() as db:
                    setting = (
                        db.query(Setting)
                        .filter(Setting.key == Setting.KEY_DOCSTORE_INTELLIGENCE_MODEL)
                        .first()
                    )
                    intelligence_model = setting.value if setting else None

            if not intelligence_model:
                logger.debug(
                    f"No intelligence model configured, skipping for store '{store.name}'"
                )
                return False

            # Sample chunks from the collection
            collection = self._chroma_client.get_collection(store.collection_name)

            # Get a diverse sample of chunks (up to 20)
            sample_size = min(20, store.chunk_count or 0)
            results = collection.get(
                limit=sample_size, include=["documents", "metadatas"]
            )

            if not results or not results.get("documents"):
                logger.warning(
                    f"No documents found in collection for store '{store.name}'"
                )
                return False

            # Build content sample for analysis
            content_samples = []
            seen_sources = set()
            for i, (doc, meta) in enumerate(
                zip(results["documents"], results["metadatas"])
            ):
                source = meta.get("source_uri", f"chunk_{i}")
                # Avoid too many chunks from same source
                source_base = source.rsplit("/", 1)[-1] if "/" in source else source
                if source_base not in seen_sources or len(seen_sources) < 5:
                    seen_sources.add(source_base)
                    # Truncate long chunks
                    content_samples.append(doc[:1500] if len(doc) > 1500 else doc)

            combined_content = "\n\n---\n\n".join(content_samples[:15])

            # Build the analysis prompt
            prompt = f"""Analyze the following document samples from a document store and generate intelligence to help with routing decisions.

DOCUMENT STORE: {store.name}
DESCRIPTION: {store.description or "(none provided)"}
SOURCE TYPE: {store.source_type}
DOCUMENT COUNT: {store.document_count}

CONTENT SAMPLES:
{combined_content[:12000]}

Based on these samples, provide:

1. THEMES: List 5-15 broad topic categories covered (as a JSON array of strings). Include both high-level domain categories (e.g., "Marketing", "Technology", "Finance") and more specific but still general sub-themes (e.g., "Semiconductor Industry", "AI Regulation", "Executive Leadership"). If content appears to be recent or time-sensitive, include a theme like "Recent News" or "Current Events". Think of these as tags that help match queries to this store.
2. BEST_FOR: Describe what types of questions this store can answer (1-2 sentences). Focus on the category of information and recency. Example: "Recent news and analysis about the technology industry, particularly semiconductor and AI developments" not specific article details.
3. SUMMARY: A brief description of the content type and coverage (1 sentence). Mention if content is recent/current. Example: "Recent industry news articles and analysis reports" not detailed descriptions of specific content.

Avoid mentioning specific company names, people, or exact dates from the documents - keep themes general enough to match related queries.

Respond in this exact format:
THEMES: ["Category1", "Category2", "Category3", ...]
BEST_FOR: <use case description>
SUMMARY: <content type description>"""

            # Call the LLM
            from providers import registry

            logger.info(
                f"Generating intelligence for store '{store.name}' using model: {intelligence_model}"
            )
            resolved = registry._resolve_actual_model(intelligence_model)
            logger.info(
                f"Resolved to provider: {resolved.provider.name}, model_id: {resolved.model_id}"
            )

            try:
                result = resolved.provider.chat_completion(
                    model=resolved.model_id,
                    messages=[{"role": "user", "content": prompt}],
                    system=None,
                    options={"max_tokens": 600, "temperature": 0.3},
                )
            except Exception as llm_error:
                logger.error(
                    f"LLM call failed for intelligence generation: {llm_error}"
                )
                raise

            response_text = result.get("content", "").strip()

            # Parse the response
            themes = []
            best_for = None
            content_summary = None

            import json
            import re

            # Extract THEMES
            themes_match = re.search(r"THEMES:\s*(\[.*?\])", response_text, re.DOTALL)
            if themes_match:
                try:
                    themes = json.loads(themes_match.group(1))
                except json.JSONDecodeError:
                    # Try to extract as comma-separated
                    themes_str = themes_match.group(1).strip("[]")
                    themes = [t.strip().strip("\"'") for t in themes_str.split(",")]

            # Extract BEST_FOR
            best_for_match = re.search(
                r"BEST_FOR:\s*(.+?)(?=\nSUMMARY:|$)", response_text, re.DOTALL
            )
            if best_for_match:
                best_for = best_for_match.group(1).strip()

            # Extract SUMMARY
            summary_match = re.search(r"SUMMARY:\s*(.+?)$", response_text, re.DOTALL)
            if summary_match:
                content_summary = summary_match.group(1).strip()

            # Save to database
            with get_db_context() as db:
                db_store = (
                    db.query(DocumentStore).filter(DocumentStore.id == store_id).first()
                )
                if db_store:
                    db_store.themes = themes if themes else None
                    db_store.best_for = best_for
                    db_store.content_summary = content_summary
                    db_store.intelligence_updated_at = datetime.utcnow()
                    db.flush()

            logger.info(
                f"Generated intelligence for store '{store.name}': "
                f"{len(themes)} themes, best_for={'yes' if best_for else 'no'}, "
                f"summary={'yes' if content_summary else 'no'}"
            )
            return True

        except Exception as e:
            logger.error(f"Failed to generate intelligence for store {store_id}: {e}")
            return False

    def start(self):
        """Start the indexer service and scheduler."""
        if self._running:
            return

        self._running = True

        # Initialize ChromaDB client
        chroma_url = os.environ.get("CHROMA_URL")
        if chroma_url:
            try:
                import chromadb

                self._chroma_client = chromadb.HttpClient(
                    host=chroma_url.replace("http://", "")
                    .replace("https://", "")
                    .split(":")[0],
                    port=int(chroma_url.split(":")[-1])
                    if ":" in chroma_url.split("/")[-1]
                    else 8000,
                )
                logger.info(f"RAG Indexer connected to ChromaDB at {chroma_url}")
            except Exception as e:
                logger.error(f"Failed to connect to ChromaDB: {e}")
                self._chroma_client = None
        else:
            logger.warning("CHROMA_URL not set - RAG indexing disabled")

        # Reset any stuck "indexing" statuses from previous runs
        # (e.g., if container was restarted mid-indexing)
        self._reset_stuck_indexing_jobs()

        # Initialize scheduler only if there are stores with schedules
        # This avoids CPU overhead from APScheduler when not needed
        from db import get_stores_with_schedule

        stores_with_schedule = get_stores_with_schedule()
        if not stores_with_schedule:
            logger.info(
                "RAG Indexer scheduler not started - no document stores have schedules"
            )
            self._scheduler = None
            return

        try:
            from apscheduler.schedulers.background import BackgroundScheduler

            self._scheduler = BackgroundScheduler()
            self._scheduler.start()
            logger.info(
                f"RAG Indexer scheduler started for {len(stores_with_schedule)} store(s)"
            )

            # Schedule existing document stores
            for store in stores_with_schedule:
                self.schedule_store(store.id, store.index_schedule)
        except ImportError as e:
            logger.warning(
                f"APScheduler not installed - scheduled indexing disabled: {e}"
            )
            self._scheduler = None
        except Exception as e:
            logger.error(f"Failed to initialize scheduler: {e}")
            self._scheduler = None

    def stop(self):
        """Stop the indexer service."""
        self._running = False

        if self._scheduler:
            self._scheduler.shutdown(wait=False)
            self._scheduler = None

        # Wait for any running indexing jobs
        for thread in self._indexing_jobs.values():
            thread.join(timeout=5)

        logger.info("RAG Indexer stopped")

    def _reset_stuck_indexing_jobs(self):
        """
        Reset any document stores or RAGs stuck in 'indexing' status.

        This handles cases where the indexing thread was killed mid-operation
        (e.g., container restart, crash) and the status was never updated.
        On startup, we reset these to 'pending' so they can be re-indexed.
        """
        from db import (
            get_all_document_stores,
            update_document_store_index_status,
        )

        try:
            stores = get_all_document_stores()
            reset_count = 0

            for store in stores:
                if store.index_status == "indexing":
                    logger.warning(
                        f"Document store '{store.name}' (ID: {store.id}) was stuck in "
                        f"'indexing' status - resetting to 'pending'"
                    )
                    update_document_store_index_status(store.id, "pending")
                    reset_count += 1

            if reset_count > 0:
                logger.info(f"Reset {reset_count} stuck indexing job(s) to 'pending'")

        except Exception as e:
            logger.error(f"Failed to reset stuck indexing jobs: {e}")

    def _schedule_all_stores(self):
        """Schedule indexing for all document stores with schedules."""
        from db import get_stores_with_schedule

        stores = get_stores_with_schedule()
        for store in stores:
            self.schedule_store(store.id, store.index_schedule)

    def schedule_store(self, store_id: int, cron_expression: str):
        """
        Schedule periodic indexing for a document store.

        Args:
            store_id: ID of the DocumentStore
            cron_expression: Cron expression (e.g., "0 2 * * *" for 2 AM daily)

        Note:
            When minute is "0", a stagger offset is added based on store_id
            to prevent all stores from indexing simultaneously (thundering herd).
            The offset is deterministic: same store always gets same minute.
        """
        if not self._scheduler:
            logger.warning("Scheduler not available - cannot schedule store indexing")
            return

        job_id = f"store_index_{store_id}"

        # Remove existing job if any
        try:
            self._scheduler.remove_job(job_id)
        except Exception:
            pass

        if not cron_expression:
            return

        try:
            parts = cron_expression.split()
            if len(parts) != 5:
                logger.error(f"Invalid cron expression: {cron_expression}")
                return

            minute = parts[0]
            hour = parts[1]
            second = 0

            # Stagger jobs that are scheduled at minute 0 to prevent thundering herd
            # Use store_id to generate deterministic offset across full hour (0-3599 seconds)
            # This supports up to 3600 jobs without collision
            if minute == "0":
                total_offset = store_id % 3600  # seconds into the hour
                stagger_minute = total_offset // 60
                second = total_offset % 60
                minute = str(stagger_minute)
                logger.debug(
                    f"Staggering store {store_id} to {stagger_minute}:{second:02d}"
                )

            from apscheduler.triggers.cron import CronTrigger

            trigger = CronTrigger(
                second=second,
                minute=minute,
                hour=hour,
                day=parts[2],
                month=parts[3],
                day_of_week=parts[4],
            )

            self._scheduler.add_job(
                self._run_store_index_job,
                trigger,
                args=[store_id],
                id=job_id,
                replace_existing=True,
            )
            actual_cron = f"{minute} {hour} {parts[2]} {parts[3]} {parts[4]}"
            logger.info(f"Scheduled indexing for store {store_id}: {actual_cron}")
        except Exception as e:
            logger.error(f"Failed to schedule store {store_id}: {e}")

    def unschedule_store(self, store_id: int):
        """Remove scheduled indexing for a document store."""
        if not self._scheduler:
            return

        job_id = f"store_index_{store_id}"
        try:
            self._scheduler.remove_job(job_id)
            logger.info(f"Unscheduled indexing for store {store_id}")
        except Exception:
            pass

    def _run_store_index_job(self, store_id: int):
        """Run document store indexing job (called by scheduler)."""
        self.index_store(store_id, background=True)

    def cancel_store_indexing(self, store_id: int) -> bool:
        """
        Cancel and reset a stuck document store indexing job.
        """
        from db import update_document_store_index_status

        job_key = f"store_{store_id}"
        with self._lock:
            # Mark as cancelled so the indexing loop will stop
            self._cancelled_jobs.add(job_key)
            if job_key in self._indexing_jobs:
                del self._indexing_jobs[job_key]

        # Reset status to pending
        update_document_store_index_status(store_id, "pending")
        logger.info(f"Cancelled/reset indexing for store {store_id}")
        return True

    def is_store_indexing(self, store_id: int) -> bool:
        """Check if a document store is currently being indexed."""
        job_key = f"store_{store_id}"
        with self._lock:
            if job_key in self._indexing_jobs:
                return self._indexing_jobs[job_key].is_alive()
        return False

    def index_store(self, store_id: int, background: bool = False) -> bool:
        """
        Index all documents for a DocumentStore.

        Args:
            store_id: ID of the DocumentStore to index
            background: If True, run in background thread

        Returns:
            True if indexing started/completed successfully
        """
        if not self._chroma_client:
            logger.error("ChromaDB not available - cannot index")
            return False

        job_key = f"store_{store_id}"

        if background:
            with self._lock:
                # Clear any previous cancellation flag
                self._cancelled_jobs.discard(job_key)

                if job_key in self._indexing_jobs:
                    thread = self._indexing_jobs[job_key]
                    if thread.is_alive():
                        logger.warning(f"Store {store_id} is already being indexed")
                        return False

                thread = threading.Thread(
                    target=self._index_store_impl,
                    args=[store_id],
                    daemon=True,
                )
                self._indexing_jobs[job_key] = thread
                thread.start()
            return True
        else:
            # Clear any previous cancellation flag for sync indexing too
            with self._lock:
                self._cancelled_jobs.discard(job_key)
            return self._index_store_impl(store_id)

    def _index_store_impl(self, store_id: int) -> bool:
        """Implementation of DocumentStore indexing."""
        from db import get_document_store_by_id, update_document_store_index_status

        from .retriever import _get_embedding_provider

        embedding_provider = None  # Initialize for cleanup in finally block

        # Get store config
        store = get_document_store_by_id(store_id)
        if not store:
            logger.error(f"Store {store_id} not found")
            return False

        if store.source_type == "local":
            source_desc = store.source_path
        elif store.source_type == "paperless":
            source_desc = f"Paperless:{os.environ.get('PAPERLESS_URL', 'unknown')}"
        elif store.source_type == "mcp:github":
            source_desc = f"GitHub:{store.github_repo or 'unknown'}"
        elif store.source_type == "website":
            source_desc = f"Website:{store.website_url or 'unknown'}"
        else:
            source_desc = f"MCP:{store.mcp_server_config.get('name', 'unknown') if store.mcp_server_config else 'unknown'}"
        logger.info(f"Starting indexing for store '{store.name}' from {source_desc}")

        # Update status to indexing
        update_document_store_index_status(store_id, "indexing")

        try:
            # Always load global vision settings (needed for local file processing)
            from .vision import get_vision_config_from_settings

            global_vision = get_vision_config_from_settings()
            vision_provider = global_vision.provider_type or "local"
            vision_model = global_vision.model_name
            vision_ollama_url = global_vision.base_url

            # Get unified source plugin for this store type
            # All source types must have a unified source plugin
            source = self._get_unified_source_for_store(store)

            if not source:
                raise ValueError(
                    f"No unified source plugin found for source type '{store.source_type}'. "
                    f"All document sources must have a plugin in builtin_plugins/unified_sources/."
                )

            if not source.is_available():
                raise ValueError(f"Document source not available: {source_desc}")

            # Get embedding provider
            embedding_provider = _get_embedding_provider(
                store.embedding_provider,
                store.embedding_model,
                store.ollama_url,
            )

            if not embedding_provider.is_available():
                raise ValueError(
                    f"Embedding provider '{store.embedding_provider}' is not available"
                )

            # Get or create ChromaDB collection
            collection_name = store.collection_name or f"docstore_{store_id}"
            collection = self._chroma_client.get_or_create_collection(
                name=collection_name,
                metadata={"hnsw:space": "cosine"},
            )

            # Get existing indexed documents for incremental indexing
            existing_docs = self._get_existing_doc_metadata(collection)
            logger.info(f"Found {len(existing_docs)} existing documents in index")

            # Find and process documents
            documents = []
            doc_count = 0
            skipped_count = 0
            deleted_uris = set()
            existing_chunk_count = (
                collection.count()
            )  # Track existing chunks for progress

            # Collect documents from source, updating progress during listing
            all_docs = []
            listing_count = 0
            for doc_info in source.list_documents():
                all_docs.append(doc_info)
                listing_count += 1
                # Update progress every 10 documents during listing phase
                if listing_count % 10 == 0:
                    update_document_store_index_status(
                        store_id,
                        "indexing",
                        document_count=listing_count,
                        chunk_count=0,
                    )
                    logger.debug(f"Listed {listing_count} documents so far...")

            total_files = len(all_docs)
            logger.info(f"Found {total_files} documents in source")

            # Apply max_documents limit if configured
            if store.max_documents and store.max_documents > 0:
                if total_files > store.max_documents:
                    logger.info(
                        f"Limiting to {store.max_documents} documents (found {total_files})"
                    )
                    all_docs = all_docs[: store.max_documents]
                    total_files = len(all_docs)

            # Track which URIs we see in source (for deletion detection)
            source_uris = {doc_info.uri for doc_info in all_docs}

            job_key = f"store_{store_id}"
            for file_idx, doc_info in enumerate(all_docs):
                # Check if cancelled
                if job_key in self._cancelled_jobs:
                    logger.info(f"Indexing cancelled for store {store_id}")
                    with self._lock:
                        self._cancelled_jobs.discard(job_key)
                    return False

                # Check if document needs re-indexing (incremental)
                if doc_info.uri in existing_docs:
                    existing_modified = existing_docs[doc_info.uri]

                    # If source provides modified_time, compare timestamps
                    if doc_info.modified_time:
                        if (
                            existing_modified
                            and existing_modified >= doc_info.modified_time
                        ):
                            # Document hasn't changed, skip
                            skipped_count += 1
                            logger.debug(
                                f"Skipping unchanged document {file_idx + 1}/{total_files}: {doc_info.name}"
                            )
                            continue
                        # Document changed - delete old chunks before re-indexing
                        self._delete_doc_chunks(collection, doc_info.uri)
                    else:
                        # No modified_time from source (e.g., Gmail) - assume immutable, skip
                        skipped_count += 1
                        logger.debug(
                            f"Skipping existing document {file_idx + 1}/{total_files}: {doc_info.name}"
                        )
                        continue

                logger.info(
                    f"Processing document {file_idx + 1}/{total_files}: {doc_info.name}"
                )
                try:
                    if store.source_type == "local":
                        # Handle both file:// URIs and plain paths
                        file_path = doc_info.uri
                        if file_path.startswith("file://"):
                            file_path = file_path[7:]  # Strip file:// prefix
                        chunks = self._process_document(
                            Path(file_path),
                            store.chunk_size,
                            store.chunk_overlap,
                            vision_provider=vision_provider,
                            vision_model=vision_model,
                            vision_ollama_url=vision_ollama_url,
                        )
                    else:
                        content = source.read_document(doc_info.uri)
                        if content:
                            chunks = self._process_content(
                                content,
                                store.chunk_size,
                                store.chunk_overlap,
                            )
                        else:
                            chunks = []

                    for chunk in chunks:
                        chunk["source_file"] = doc_info.name
                        chunk["source_uri"] = doc_info.uri
                        chunk["modified_time"] = (
                            doc_info.modified_time
                        )  # Track for incremental
                    documents.extend(chunks)
                    doc_count += 1
                    logger.info(
                        f"Completed {doc_info.name}: {len(chunks)} chunks "
                        f"(total: {len(documents)} chunks from {doc_count} docs, skipped: {skipped_count})"
                    )

                    update_document_store_index_status(
                        store_id,
                        "indexing",
                        document_count=total_files,
                        chunk_count=existing_chunk_count + len(documents),
                    )
                except Exception as e:
                    logger.warning(f"Failed to process {doc_info.uri}: {e}")

            # Delete documents that no longer exist in source
            for uri in existing_docs:
                if uri not in source_uris:
                    self._delete_doc_chunks(collection, uri)
                    deleted_uris.add(uri)

            if deleted_uris:
                logger.info(
                    f"Deleted {len(deleted_uris)} documents no longer in source"
                )

            if not documents:
                if skipped_count > 0:
                    # All documents were unchanged - still a success
                    logger.info(
                        f"No new/modified documents to index (skipped {skipped_count} unchanged)"
                    )
                    # Get current chunk count from collection
                    existing_chunk_count = collection.count()
                    update_document_store_index_status(
                        store_id,
                        "ready",
                        document_count=skipped_count,
                        chunk_count=existing_chunk_count,
                        collection_name=collection_name,
                    )
                else:
                    logger.warning(f"No documents found in {source_desc}")
                    update_document_store_index_status(
                        store_id,
                        "ready",
                        document_count=0,
                        chunk_count=0,
                        collection_name=collection_name,
                    )
                return True

            # Generate embeddings and store in ChromaDB
            logger.info(
                f"Generating embeddings for {len(documents)} chunks "
                f"({doc_count} new/modified docs, {skipped_count} unchanged)..."
            )

            batch_size = 100
            total_chunks = 0

            for i in range(0, len(documents), batch_size):
                # Check if cancelled during embedding
                if job_key in self._cancelled_jobs:
                    logger.info(
                        f"Indexing cancelled during embedding for store {store_id}"
                    )
                    with self._lock:
                        self._cancelled_jobs.discard(job_key)
                    return False

                batch = documents[i : i + batch_size]
                texts = [doc["content"] for doc in batch]

                embed_result = embedding_provider.embed(texts)

                ids = [
                    hashlib.md5(
                        f"{doc.get('source_uri', doc['source_file'])}:{doc['chunk_index']}".encode()
                    ).hexdigest()
                    for doc in batch
                ]
                metadatas = []
                for doc in batch:
                    meta = {
                        "source_file": doc["source_file"],
                        "source_uri": doc.get("source_uri", doc["source_file"]),
                        "chunk_index": doc["chunk_index"],
                        "indexed_at": datetime.utcnow().isoformat(),
                        "store_id": store_id,
                        "store_name": store.name,
                    }
                    # Add modified_time for incremental indexing
                    if doc.get("modified_time"):
                        meta["modified_time"] = doc["modified_time"]
                        # Add document_date (YYYY-MM-DD) for temporal filtering
                        # This enables date-based queries like "news from last week"
                        try:
                            # Parse ISO datetime and extract date
                            doc_date = doc["modified_time"][:10]  # YYYY-MM-DD
                            meta["document_date"] = doc_date
                        except (TypeError, IndexError):
                            pass
                    # Add document-level metadata (dates, location, etc.)
                    if doc.get("metadata"):
                        for key, value in doc["metadata"].items():
                            if value is not None:
                                # ChromaDB only accepts scalar metadata values
                                # Convert lists to comma-separated strings
                                if isinstance(value, list):
                                    meta[key] = ", ".join(str(v) for v in value)
                                elif isinstance(value, (str, int, float, bool)):
                                    meta[key] = value
                                # Skip other non-scalar types
                    metadatas.append(meta)

                collection.add(
                    ids=ids,
                    embeddings=embed_result.embeddings,
                    documents=texts,
                    metadatas=metadatas,
                )

                total_chunks += len(batch)
                logger.debug(f"Indexed {total_chunks}/{len(documents)} chunks")

            # Get total chunk count (new + existing unchanged)
            final_chunk_count = collection.count()
            total_doc_count = doc_count + skipped_count

            update_document_store_index_status(
                store_id,
                "ready",
                document_count=total_doc_count,
                chunk_count=final_chunk_count,
                collection_name=collection_name,
            )

            logger.info(
                f"Indexing complete for store '{store.name}': "
                f"{total_doc_count} documents ({doc_count} new/modified, {skipped_count} unchanged), "
                f"{final_chunk_count} total chunks"
            )

            # Generate intelligence for the store (themes, best_for, summary)
            # This runs after indexing to analyze the content
            try:
                self.generate_store_intelligence(store_id)
            except Exception as e:
                logger.warning(
                    f"Intelligence generation failed for store '{store.name}': {e}"
                )
            return True

        except Exception as e:
            logger.error(f"Indexing failed for store {store_id}: {e}")
            update_document_store_index_status(store_id, "error", error=str(e))
            return False

        finally:
            # Clean up GPU memory
            if embedding_provider and hasattr(embedding_provider, "unload"):
                embedding_provider.unload()
            self._clear_gpu_memory()

            job_key = f"store_{store_id}"
            with self._lock:
                self._indexing_jobs.pop(job_key, None)

    def _find_documents(self, source_path: Path):
        """Find all supported documents in a directory."""
        for file_path in source_path.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in SUPPORTED_EXTENSIONS:
                yield file_path

    def _process_document(
        self,
        file_path: Path,
        chunk_size: int,
        chunk_overlap: int,
        vision_provider: str = "local",
        vision_model: Optional[str] = None,
        vision_ollama_url: Optional[str] = None,
    ) -> list[dict]:
        """
        Process a document and return chunks.

        Uses configured PDF parser for PDFs, Docling for other formats.

        Args:
            file_path: Path to the document
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks
            vision_provider: Vision model provider ("local", "ollama:<instance>", or provider name)
            vision_model: Vision model name for remote processing
            vision_ollama_url: Ollama URL for vision processing
        """
        suffix = file_path.suffix.lower()

        # For PDFs, check the configured parser setting
        if suffix == ".pdf":
            pdf_parser = self._get_pdf_parser_setting()

            if pdf_parser == "pypdf":
                try:
                    return self._process_pdf_with_pypdf(
                        file_path, chunk_size, chunk_overlap
                    )
                except Exception:
                    # Fall through to Docling
                    pass
            elif pdf_parser == "jina":
                # Jina doesn't support local files, fall through to Docling
                logger.info(
                    f"Jina parser not supported for local files ({file_path.name}), "
                    "using Docling"
                )
            # Fall through to Docling for PDFs

        # Try Docling (for PDFs with docling setting, and all other document types)
        try:
            return self._process_with_docling(
                file_path,
                chunk_size,
                chunk_overlap,
                vision_provider=vision_provider,
                vision_model=vision_model,
                vision_ollama_url=vision_ollama_url,
            )
        except ImportError:
            logger.debug("Docling not available, using fallback parser")
        except Exception as e:
            logger.warning(f"Docling failed for {file_path}: {e}, trying fallback")

        # Fallback to simple parsing
        return self._process_simple(file_path, suffix, chunk_size, chunk_overlap)

    def _get_pdf_parser_setting(self) -> str:
        """Get the configured RAG PDF parser from settings."""
        try:
            from db.connection import get_db_context
            from db.models import Setting

            with get_db_context() as db:
                setting = (
                    db.query(Setting)
                    .filter(Setting.key == Setting.KEY_RAG_PDF_PARSER)
                    .first()
                )
                return setting.value if setting else "docling"
        except Exception:
            return "docling"

    def _process_pdf_with_pypdf(
        self, file_path: Path, chunk_size: int, chunk_overlap: int
    ) -> list[dict]:
        """Process PDF using pypdf (fast, text-only)."""
        try:
            import pypdf

            reader = pypdf.PdfReader(str(file_path))
            text_parts = []
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)

            text = "\n\n".join(text_parts)

            if not text.strip():
                logger.warning(
                    f"pypdf extracted no text from {file_path} (possibly scanned), "
                    "falling back to Docling"
                )
                raise ValueError("No text extracted")

            logger.info(f"Processed {file_path.name} with pypdf")
            return self._chunk_text(text, chunk_size, chunk_overlap)

        except Exception as e:
            logger.warning(
                f"pypdf failed for {file_path}: {e}, falling back to Docling"
            )
            raise

    def _process_pdf_with_jina(
        self, file_path: Path, chunk_size: int, chunk_overlap: int
    ) -> list[dict]:
        """Process PDF using Jina Reader API - NOT supported for local files."""
        # Jina Reader requires a URL, not local files
        # This should not be called for local document stores
        logger.info(
            f"Jina parser not supported for local files ({file_path.name}), "
            "falling back to Docling"
        )
        raise ValueError("Jina not supported for local files")

    def _process_with_docling(
        self,
        file_path: Path,
        chunk_size: int,
        chunk_overlap: int,
        vision_provider: str = "local",
        vision_model: Optional[str] = None,
        vision_ollama_url: Optional[str] = None,
    ) -> list[dict]:
        """Process document using Docling with configurable vision model."""
        from .vision import get_document_converter

        converter = get_document_converter(
            vision_provider=vision_provider,
            vision_model=vision_model,
            vision_ollama_url=vision_ollama_url,
        )
        try:
            result = converter.convert(str(file_path))

            # Get text content
            text = result.document.export_to_markdown()

            # Chunk the text
            return self._chunk_text(text, chunk_size, chunk_overlap)
        finally:
            # Clean up converter and GPU memory after each document
            # to prevent memory accumulation during batch processing
            del converter
            self._clear_gpu_memory()

    def _process_simple(
        self, file_path: Path, suffix: str, chunk_size: int, chunk_overlap: int
    ) -> list[dict]:
        """Simple fallback document processing."""
        text = ""

        if suffix in {".txt", ".md", ".asciidoc", ".adoc"}:
            text = file_path.read_text(encoding="utf-8", errors="ignore")

        elif suffix in {".html", ".htm"}:
            try:
                from bs4 import BeautifulSoup

                html = file_path.read_text(encoding="utf-8", errors="ignore")
                soup = BeautifulSoup(html, "html.parser")
                text = soup.get_text(separator="\n")
            except ImportError:
                # Just strip tags crudely
                import re

                html = file_path.read_text(encoding="utf-8", errors="ignore")
                text = re.sub(r"<[^>]+>", "", html)

        elif suffix == ".pdf":
            try:
                import pypdf

                reader = pypdf.PdfReader(str(file_path))
                text = "\n".join(page.extract_text() or "" for page in reader.pages)
            except ImportError:
                raise ValueError("pypdf required for PDF processing without Docling")

        elif suffix == ".docx":
            try:
                from docx import Document

                doc = Document(str(file_path))
                text = "\n".join(para.text for para in doc.paragraphs)
            except ImportError:
                raise ValueError(
                    "python-docx required for DOCX processing without Docling"
                )

        else:
            raise ValueError(f"Unsupported file type: {suffix}")

        if not text.strip():
            return []

        return self._chunk_text(text, chunk_size, chunk_overlap)

    def _process_content(
        self,
        content: "DocumentContent",
        chunk_size: int,
        chunk_overlap: int,
    ) -> list[dict]:
        """
        Process document content from an MCP source.

        Handles both text and binary content.

        Args:
            content: DocumentContent from MCP source
            chunk_size: Target chunk size in characters
            chunk_overlap: Overlap between chunks

        Returns:
            List of chunk dicts with 'content', 'chunk_index', and optional metadata
        """
        from plugin_base.document_source import DocumentContent

        # If we have text content, chunk it directly
        if content.text:
            text = content.text

            # Handle HTML content
            if content.mime_type and "html" in content.mime_type.lower():
                try:
                    from bs4 import BeautifulSoup

                    soup = BeautifulSoup(text, "html.parser")
                    text = soup.get_text(separator="\n")
                except ImportError:
                    import re

                    text = re.sub(r"<[^>]+>", "", text)

            chunks = self._chunk_text(text, chunk_size, chunk_overlap)

            # Add metadata from DocumentContent to each chunk
            if content.metadata:
                for chunk in chunks:
                    chunk["metadata"] = content.metadata

            return chunks

        # If we have binary content, try to extract text
        if content.binary:
            mime = content.mime_type or ""

            # PDF
            if "pdf" in mime.lower():
                try:
                    import io

                    import pypdf

                    reader = pypdf.PdfReader(io.BytesIO(content.binary))
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    return self._chunk_text(text, chunk_size, chunk_overlap)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF: {e}")
                    return []

            # DOCX
            if "wordprocessingml" in mime.lower() or content.metadata.get("filename", "").endswith(".docx"):
                try:
                    import io

                    from docx import Document

                    doc = Document(io.BytesIO(content.binary))
                    text = "\n".join(para.text for para in doc.paragraphs)
                    return self._chunk_text(text, chunk_size, chunk_overlap)
                except Exception as e:
                    logger.warning(f"Failed to extract text from DOCX: {e}")
                    return []

            # For images and other binary formats, we can't extract text
            # without vision models - skip for now
            logger.debug(f"Skipping binary content with mime type: {mime}")
            return []

        return []

    def _chunk_text(self, text: str, chunk_size: int, chunk_overlap: int) -> list[dict]:
        """
        Split text into overlapping chunks.

        Simple character-based chunking with sentence awareness.
        """
        if not text.strip():
            return []

        # Normalize whitespace
        text = " ".join(text.split())

        chunks = []
        start = 0
        chunk_index = 0

        while start < len(text):
            # Calculate end position
            end = start + chunk_size

            # Try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings
                for sep in [". ", ".\n", "! ", "? ", "\n\n"]:
                    last_sep = text.rfind(sep, start, end)
                    if last_sep > start + chunk_size // 2:
                        end = last_sep + len(sep)
                        break

            chunk_text = text[start:end].strip()
            if chunk_text:
                chunks.append(
                    {
                        "content": chunk_text,
                        "chunk_index": chunk_index,
                    }
                )
                chunk_index += 1

            # Move start with overlap
            start = end - chunk_overlap
            if start >= len(text) - chunk_overlap:
                break

        return chunks

    def delete_collection(self, rag_id: int) -> bool:
        """Delete the ChromaDB collection for a legacy RAG."""
        if not self._chroma_client:
            return False

        from db import get_smart_rag_by_id

        rag = get_smart_rag_by_id(rag_id)
        if not rag or not rag.collection_name:
            return False

        try:
            self._chroma_client.delete_collection(rag.collection_name)
            logger.info(f"Deleted collection {rag.collection_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to delete collection: {e}")
            return False

    def delete_store_collection(self, store_id: int) -> bool:
        """Delete the ChromaDB collection for a document store."""
        if not self._chroma_client:
            return False

        from db import get_document_store_by_id

        store = get_document_store_by_id(store_id)
        if not store or not store.collection_name:
            return False

        try:
            self._chroma_client.delete_collection(store.collection_name)
            logger.info(f"Deleted collection {store.collection_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to delete collection: {e}")
            return False

    def get_collection_stats(self, rag_id: int) -> Optional[dict]:
        """Get statistics for a legacy RAG's ChromaDB collection."""
        if not self._chroma_client:
            return None

        from db import get_smart_rag_by_id

        rag = get_smart_rag_by_id(rag_id)
        if not rag or not rag.collection_name:
            return None

        try:
            collection = self._chroma_client.get_collection(rag.collection_name)
            count = collection.count()
            return {
                "collection_name": rag.collection_name,
                "chunk_count": count,
            }
        except Exception:
            return None

    def get_store_collection_stats(self, store_id: int) -> Optional[dict]:
        """Get statistics for a document store's ChromaDB collection."""
        if not self._chroma_client:
            return None

        from db import get_document_store_by_id

        store = get_document_store_by_id(store_id)
        if not store or not store.collection_name:
            return None

        try:
            collection = self._chroma_client.get_collection(store.collection_name)
            count = collection.count()
            return {
                "collection_name": store.collection_name,
                "chunk_count": count,
            }
        except Exception:
            return None


# Global indexer instance
_indexer: Optional[RAGIndexer] = None


def get_indexer() -> RAGIndexer:
    """Get or create the global RAG indexer instance."""
    global _indexer
    if _indexer is None:
        _indexer = RAGIndexer()
        _indexer.start()
    return _indexer


def start_indexer():
    """Start the global RAG indexer."""
    indexer = get_indexer()
    indexer.start()


def stop_indexer():
    """Stop the global RAG indexer."""
    global _indexer
    if _indexer:
        _indexer.stop()
        _indexer = None
